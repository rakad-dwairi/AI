from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = PROJECT_ROOT / "phase1_validation" / "reports" / "sample_01_speaker_evaluation.json"
OUTPUT_PATH = (
    PROJECT_ROOT
    / "phase1_validation"
    / "reports"
    / "Thohor_Stakeholder_Evaluation_Report_sample_01.docx"
)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 102, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF6EF"
PALE_GOLD = "FFF4CC"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"


def main() -> None:
    data = json.loads(REPORT_PATH.read_text(encoding="utf-8"))
    rows = data["criteria"]
    doc = Document()
    setup_document(doc)

    add_title(doc, data)
    add_executive_callout(doc, data, rows)
    add_workflow(doc)
    add_key_metrics(doc, rows)
    add_results_table(doc, rows)
    add_arabic_summary(doc, data, rows)
    add_notes(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Thohor validation report | generated from sample_01")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_title(doc: Document, data: dict[str, Any]) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Thohor Video Evaluation Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run(
        "Stakeholder brief based on استمارة 2026.xlsx and the current AI/API evaluation pipeline"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    meta = [
        ("Sample", data["sample_id"]),
        ("Reference rubric", data["source_reference"]),
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Report source", "speaker_evaluation.json"),
    ]
    table = doc.add_table(rows=1, cols=4)
    style_table(table, [1.15, 1.65, 1.35, 2.35], header=False, fill=LIGHT_GRAY)
    row = table.rows[0]
    for idx, (label, value) in enumerate(meta):
        cell = row.cells[idx]
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        a = p.add_run(label + "\n")
        a.bold = True
        set_run_font(a, 8.5, DARK_BLUE)
        b = p.add_run(str(value))
        set_run_font(b, 8.5, INK)


def add_executive_callout(doc: Document, data: dict[str, Any], rows: list[dict[str, Any]]) -> None:
    summary = data["summary"]
    scored = summary["scored"]
    total = summary["total_criteria"]
    avg = summary["average_score_for_scored_criteria"]
    human = summary["human_review_required"]
    high = sum(1 for row in rows if safe_score(row) >= 80)
    weak = sum(1 for row in rows if row.get("rating") == "weak")

    doc.add_heading("Executive Snapshot", level=1)
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.5], header=False, fill=PALE_GREEN)
    cell = table.cell(0, 0)
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        f"The pipeline evaluated {total} Excel criteria for sample_01. "
        f"{scored} criteria received automatic scores, with an average scored result of {avg}. "
        f"{high} scored criteria reached very_good or excellent, while {weak} scored criteria were weak. "
        f"{human} criteria remain marked as human review because the Excel form requires expert judgment."
    )
    set_run_font(r, 10.5, INK)


def add_workflow(doc: Document) -> None:
    doc.add_heading("What We Built and Ran", level=1)
    steps = [
        (
            "1",
            "Read the evaluation reference",
            "Loaded the criteria and evaluation methods from استمارة 2026.xlsx. The tool uses the Excel file as the source of truth.",
        ),
        (
            "2",
            "Prepared the video sample",
            "Registered VideoOne as sample_01, extracted the audio track, and sampled video frames for visual analysis.",
        ),
        (
            "3",
            "Ran the analysis tools",
            "Used Deepgram/OpenAI for speech and transcript signals, local audio for volume/pitch/silence, MediaPipe for body and hand signals, and AWS Rekognition for face/head-pose signals.",
        ),
        (
            "4",
            "Normalized the outputs",
            "Converted raw API/tool outputs into comparable factor signals such as words_per_minute, pause_rate, hand_movement, posture, smile_rate, and duration_suitability.",
        ),
        (
            "5",
            "Mapped signals to Excel criteria",
            "Matched every criterion to the relevant signals and assigned a score only when the evidence supported it.",
        ),
        (
            "6",
            "Generated stakeholder outputs",
            "Created Markdown, JSON, CSV, and this DOCX report with score basis, evidence, and limitations.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    style_table(table, [0.45, 1.65, 4.4], header=True)
    set_header(table.rows[0].cells, ["Step", "Stage", "What happened"])
    for item in steps:
        new_row = table.add_row()
        keep_row_together(new_row)
        cells = new_row.cells
        for cell, value in zip(cells, item):
            set_cell_text(cell, value, size=9)


def add_key_metrics(doc: Document, rows: list[dict[str, Any]]) -> None:
    doc.add_heading("Key Measured Evidence", level=1)
    metrics = [
        ("Speech speed", "Row 19", get_value(rows, 19), "words per minute", "Deepgram"),
        ("Pause rate", "Row 15", get_value(rows, 15), "pauses per 100 words", "Deepgram"),
        ("Hand movement", "Row 30", get_value(rows, 30, "movement_event_count"), "movement events", "MediaPipe"),
        (
            "Hand movement rate",
            "Row 30",
            get_value(rows, 30, "movement_events_per_minute"),
            "events per minute",
            "MediaPipe",
        ),
        ("Hand visibility", "Row 30", get_value(rows, 30, "hand_visible_frames"), "visible sampled frames", "MediaPipe"),
        ("Body movement", "Row 22/23", get_supporting_value(rows, 22, "body_movement", "movement_event_count"), "movement events", "MediaPipe"),
        ("Smile rate", "Row 34", get_value(rows, 34), "smiles per minute / sampled-frame proxy", "AWS/MediaPipe normalized"),
        ("Actual duration", "Row 50", get_value(rows, 50, "actual_duration_seconds"), "seconds", "Deepgram"),
    ]
    table = doc.add_table(rows=1, cols=5)
    style_table(table, [1.35, 0.75, 1.15, 1.85, 1.4], header=True)
    set_header(table.rows[0].cells, ["Metric", "Criterion", "Value", "Meaning", "Source"])
    for metric in metrics:
        new_row = table.add_row()
        keep_row_together(new_row)
        cells = new_row.cells
        for cell, value in zip(cells, metric):
            set_cell_text(cell, clean_value(value), size=8.8)


def add_results_table(doc: Document, rows: list[dict[str, Any]]) -> None:
    doc.add_heading("Criteria Results and Evidence", level=1)
    note = doc.add_paragraph()
    note_run = note.add_run(
        "This table is a stakeholder summary of all 44 Excel criteria. Detailed machine-readable evidence is preserved in the JSON and CSV reports."
    )
    set_run_font(note_run, 9.5, INK)

    table = doc.add_table(rows=1, cols=6)
    style_table(table, [0.42, 1.05, 1.72, 0.72, 0.78, 1.81], header=True)
    set_header(table.rows[0].cells, ["Row", "Axis", "Criterion", "Score", "Rating", "Evidence basis"])
    for row_data in rows:
        new_row = table.add_row()
        keep_row_together(new_row)
        cells = new_row.cells
        values = [
            row_data["source_row"],
            row_data["axis"],
            row_data["criterion"],
            row_data["score"] if row_data["score"] != "" else "N/A",
            display_rating(row_data),
            evidence_summary(row_data),
        ]
        for i, (cell, value) in enumerate(zip(cells, values)):
            set_cell_text(cell, str(value), size=7.6 if i in (2, 5) else 8)
        shade_row(cells, row_data)


def add_arabic_summary(doc: Document, data: dict[str, Any], rows: list[dict[str, Any]]) -> None:
    doc.add_heading("Brief Arabic Stakeholder Summary", level=1)
    wpm = get_value(rows, 19)
    hand_events = get_value(rows, 30, "movement_event_count")
    hand_rate = get_value(rows, 30, "movement_events_per_minute")
    pauses = get_supporting_value(rows, 15, "pause_count", "count")
    pause_rate = get_value(rows, 15)
    duration = get_value(rows, 50, "actual_duration_seconds")
    expected = get_value(rows, 50, "expected_duration_seconds")
    posture = row_by_number(rows, 22)
    smile = row_by_number(rows, 34)
    avg = data["summary"]["average_score_for_scored_criteria"]

    arabic = (
        f"لقد تم تقييم الفيديو بناء على معايير استمارة 2026. "
        f"تحدث الشخص بسرعة تقارب {fmt(wpm)} كلمة في الدقيقة، "
        f"وتم رصد {fmt(hand_events)} حركة يد تقريبا ضمن العينة، بمعدل {fmt(hand_rate)} حركة في الدقيقة. "
        f"كما تم رصد {fmt(pauses)} وقفات مؤثرة بمعدل {fmt(pause_rate)} وقفة لكل 100 كلمة. "
        f"مدة الحديث الفعلية كانت {fmt(duration)} ثانية مقارنة بالمدة المتوقعة {fmt(expected)} ثانية. "
        f"حصل معيار لغة الجسد المنضبطة على {posture['score']} ({posture['rating']}), "
        f"وحصل معيار الابتسامة على {smile['score']} ({smile['rating']}). "
        f"المتوسط العام للمعايير التي تم تسجيلها آليا هو {fmt(avg)}. "
        "المعايير التي بقيت بدون رقم هي معايير تتطلب مراجعة بشرية حسب ملف Excel، وليست فشلا في تشغيل الأداة."
    )
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(arabic)
    set_run_font(run, 11, INK)


def add_notes(doc: Document) -> None:
    doc.add_heading("Important Interpretation Notes", level=1)
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.5], header=False, fill=PALE_GOLD)
    cell = table.cell(0, 0)
    clear_cell(cell)
    p = cell.paragraphs[0]
    note_run = p.add_run(
        "Several current scores are proxy-based. For example, pronunciation clarity uses ASR agreement rather than full phoneme alignment, gaze uses head-pose proxies, and hand movement is based on sampled-frame wrist-motion events. These are useful for Phase 1 validation, but operational deployment should refine the proxy measurements where the rubric demands stricter human-level judgment."
    )
    set_run_font(note_run, 9.5, INK)


def style_table(table, widths_in: list[float], header: bool = False, fill: str | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths_in)))
    tbl_w.set(qn("w:type"), "dxa")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D0D7DE")
        borders.append(tag)
    tbl_pr.append(borders)
    margins = OxmlElement("w:tblCellMar")
    for side, width in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        mar = OxmlElement(f"w:{side}")
        mar.set(qn("w:w"), str(width))
        mar.set(qn("w:type"), "dxa")
        margins.append(mar)
    tbl_pr.append(margins)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, Inches(widths_in[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if fill:
                shade_cell(cell, fill)
    if header:
        repeat_header(table.rows[0])


def set_header(cells, labels: list[str]) -> None:
    for cell, label in zip(cells, labels):
        shade_cell(cell, LIGHT_GRAY)
        set_cell_text(cell, label, bold=True, size=8.5, color=DARK_BLUE)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.5, color: RGBColor = INK) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.bold = bold
    set_run_font(r, size, color)
    if has_arabic(text):
        set_rtl(p)


def set_run_font(run, size: float, color: RGBColor) -> None:
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
    sz_cs = run._element.rPr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        run._element.rPr.append(sz_cs)
    sz_cs.set(qn("w:val"), str(int(size * 2)))


def has_arabic(text: str) -> bool:
    return any("\u0600" <= char <= "\u06ff" for char in text)


def clear_cell(cell) -> None:
    for p in cell.paragraphs:
        p.clear()


def shade_row(cells, row_data: dict[str, Any]) -> None:
    if row_data["evaluation_status"] == "human_review_required":
        fill = PALE_GOLD
    elif row_data.get("rating") == "weak":
        fill = PALE_RED
    elif safe_score(row_data) >= 80:
        fill = PALE_GREEN
    else:
        fill = WHITE
    for cell in cells:
        shade_cell(cell, fill)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def row_by_number(rows: list[dict[str, Any]], number: int) -> dict[str, Any]:
    return next(row for row in rows if row["source_row"] == number)


def safe_score(row: dict[str, Any]) -> float:
    try:
        return float(row.get("score") or 0)
    except (TypeError, ValueError):
        return 0.0


def selected_signal(row: dict[str, Any]) -> dict[str, Any] | None:
    basis = row.get("score_basis") or {}
    return basis.get("selected_signal")


def get_value(rows: list[dict[str, Any]], row_number: int, key: str | None = None) -> Any:
    signal = selected_signal(row_by_number(rows, row_number))
    if not signal:
        return ""
    value = signal.get("value")
    if key and isinstance(value, dict):
        return value.get(key, "")
    return value


def get_supporting_value(
    rows: list[dict[str, Any]],
    row_number: int,
    factor_id: str,
    key: str | None = None,
) -> Any:
    row = row_by_number(rows, row_number)
    for signal in row.get("detailed_evidence") or []:
        if signal.get("factor_id") == factor_id:
            value = signal.get("value")
            if key and isinstance(value, dict):
                return value.get(key, "")
            return value
    return ""


def evidence_summary(row: dict[str, Any]) -> str:
    signal = selected_signal(row)
    if not signal:
        if row["evaluation_status"] == "human_review_required":
            return "Excel marks this criterion as human review; supporting signals may exist but no automatic score was assigned."
        return row.get("explanation", "")
    factor = signal["factor_name"]
    value = signal.get("value")
    if str(signal.get("factor_id", "")).startswith("criterion_score_") and isinstance(value, dict):
        explanation = value.get("explanation") or value.get("status") or "Transcript evidence available"
        return f"Transcript rubric score. Evidence: {truncate(str(explanation), 130)} Confidence {signal.get('confidence')}."
    if isinstance(value, dict):
        compact = ", ".join(
            f"{k}: {clean_value(v)}"
            for k, v in value.items()
            if k not in {"explanation", "criterion", "source", "status", "rating"}
        )
    else:
        compact = clean_value(value)
    return f"{factor}: {truncate(compact, 150)}. Confidence {signal.get('confidence')}."


def display_rating(row: dict[str, Any]) -> str:
    value = row["rating"] if row["rating"] else row["evaluation_status"]
    return {
        "excellent": "Excellent",
        "very_good": "Very good",
        "good": "Good",
        "needs_development": "Needs dev.",
        "weak": "Weak",
        "human_review_required": "Human review",
        "partial_evidence_only": "Partial",
        "not_supported": "Not supported",
    }.get(value, str(value))


def truncate(text: str, limit: int) -> str:
    text = " ".join(text.split())
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "..."


def clean_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return fmt(value)
    return str(value)


def fmt(value: Any) -> str:
    if value == "":
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if number.is_integer():
        return str(int(number))
    return f"{number:.2f}"


if __name__ == "__main__":
    main()
